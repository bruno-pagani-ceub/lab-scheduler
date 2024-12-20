from datetime import datetime, timedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class WeeklySchedule:
    def __init__(self):
        self.doc = Document()

    def generate_schedule_doc(self, data, lab, start_of_week):
        schedule = {}
        unique_days = set()
        unique_times = set()

        for row in data:
            dt = row['dt_dia']
            hr_inicio = row['hr_inicio']
            hr_fim = row['hr_fim']

            unique_days.add(dt)
            unique_times.add((hr_inicio, hr_fim))

            key = (dt, (hr_inicio, hr_fim))
            schedule[key] = {
                'reserva_id': row['reserva_id'],
                'cod_reserva': row['cod_reserva'],
                'usuario': row['nm_usuario'],
                'ds_identificacao': row['ds_identificacao'],
                'tipo_usuario': row['ds_tipo_usuario'],
                'tp_reserva': row['tp_reserva'],
                'is_ativa': row['is_ativa']
            }

        unique_days = sorted(unique_days)
        unique_times = sorted(unique_times, key=lambda x: x[0])

       
        heading = self.doc.add_heading("Cronograma Semanal", level=0)
       
        for run in heading.runs:
            run.font.size = Pt(14)

        if data:
            first = data[0]
            lab_info = f"Lab: {first['ds_bloco']} - {first['ds_sala']} (ID: {first['lab_id']})"
        else:
            lab_info = "No data found for this lab and week."
        p = self.doc.add_paragraph(lab_info)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        for run in p.runs:
            run.font.size = Pt(10)

        all_days_in_week = [(start_of_week + timedelta(days=i)) for i in range(7)]
        
       
        table = self.doc.add_table(rows=len(unique_times) + 1, cols=8)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Medium Shading 1 Accent 1"

       
        header_cells = table.rows[0].cells
        header_cells[0].text = "Horário"
        for i, day_date in enumerate(all_days_in_week):
            header_cells[i+1].text = day_date.strftime("%a\n%d/%m")

       
        for r_idx, time_slot in enumerate(unique_times, start=1):
            time_cell = table.rows[r_idx].cells[0]
            time_cell.text = f"{time_slot[0]} - {time_slot[1]}"

            for c_idx, day_date in enumerate(all_days_in_week, start=1):
                cell = table.rows[r_idx].cells[c_idx]

                cell.text = ""

                key = (day_date.date(), time_slot)
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.add_run()

                if key in schedule:
                    info = schedule[key]
                    if info['reserva_id']:
                        run.text = f"Reservado por {info['usuario']} ({info['cod_reserva']}) - {info['tp_reserva']}"
                    else:
                        run.text = "Disponível"
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                else:
                    run.text = "Sem horário"
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) 

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        doc_name = f"weekly_schedule_{lab}_{start_of_week.strftime('%Y%m%d')}.docx"
        self.doc.save(doc_name)
        return doc_name
